from fuzzywuzzy import fuzz
import numpy as np
from docx import Document
from docx.shared import RGBColor


def combine_files():
    my_lyrics = r'my_lyrics/file.txt'
    official_lyrics = r'official_lyrics/file.txt'
    scores = []
    with open(my_lyrics, 'r') as my_lyrics_file, open(official_lyrics, 'r') as official_lyrics_file:
        my_lyrics_rows = my_lyrics_file.readlines()
        official_lyrics_rows = official_lyrics_file.readlines()

        result = Document()

        for my_lyrics_row, official_lyrics_row in zip(my_lyrics_rows, official_lyrics_rows):
            my_lyrics_row = my_lyrics_row.replace('\n', '').strip()
            official_lyrics_row = official_lyrics_row.replace('\n', '').strip()

            paragraph_1 = result.add_paragraph(my_lyrics_row)
            for run in paragraph_1.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)

            paragraph_2 = result.add_paragraph(official_lyrics_row)
            for run in paragraph_2.runs:
                run.font.color.rgb = RGBColor(0, 0, 255)

            result.add_paragraph('')

            partial = fuzz.ratio(my_lyrics_row, official_lyrics_row)
            scores.append(partial)
        result.save('results.doc')
    return np.mean(scores)


score = combine_files()
print(f'Your score was: {score}/100.0!!')
print(f'Text saved in results.doc')
